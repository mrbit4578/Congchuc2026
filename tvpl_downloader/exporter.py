import os
import re
import json
import logging
from typing import Dict, Any
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("tvpl_downloader.exporter")

class DocumentExporter:
    """Exports scraped document data into Markdown, JSON, and Word (.docx) files."""

    @staticmethod
    def export_all(scraped_data: Dict[str, Any], output_dir: str) -> Dict[str, str]:
        os.makedirs(output_dir, exist_ok=True)
        
        slug = DocumentExporter._get_slug(scraped_data["metadata"], scraped_data.get("url", ""))
        
        md_path = os.path.join(output_dir, f"{slug}.md")
        json_path = os.path.join(output_dir, f"{slug}.json")
        docx_path = os.path.join(output_dir, f"{slug}.docx")
        
        DocumentExporter.export_json(scraped_data, json_path)
        DocumentExporter.export_markdown(scraped_data, md_path)
        DocumentExporter.export_docx(scraped_data, docx_path)
        
        return {
            "markdown": md_path,
            "json": json_path,
            "docx": docx_path
        }

    @staticmethod
    def export_json(scraped_data: Dict[str, Any], file_path: str):
        data = {
            "url": scraped_data.get("url"),
            "metadata": scraped_data.get("metadata"),
            "text_content": scraped_data.get("text_content"),
            "downloaded_files": scraped_data.get("downloaded_files", [])
        }
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"Exported JSON: {file_path}")

    @staticmethod
    def export_markdown(scraped_data: Dict[str, Any], file_path: str):
        meta = scraped_data.get("metadata", {})
        title = meta.get("title", "Văn bản pháp luật")
        
        lines = [
            f"# {title}",
            "",
            "## 📋 THÔNG TIN VĂN BẢN",
            f"- **Số hiệu:** {meta.get('so_hieu', 'N/A')}",
            f"- **Loại văn bản:** {meta.get('loai_van_ban', 'N/A')}",
            f"- **Ngày ban hành:** {meta.get('ngay_ban_hanh', 'N/A')}",
            f"- **Ngày hiệu lực:** {meta.get('ngay_hieu_luc', 'N/A')}",
            f"- **Cơ quan ban hành:** {meta.get('co_quan_ban_hanh', 'N/A')}",
            f"- **Tình trạng:** {meta.get('tinh_trang', 'N/A')}",
            f"- **Nguồn:** [{scraped_data.get('url')}]({scraped_data.get('url')})",
            "",
            "---",
            "",
            "## 📜 NỘI DUNG VĂN BẢN",
            "",
        ]
        
        html_content = scraped_data.get("html_content", "")
        if html_content:
            soup = BeautifulSoup(html_content, "lxml")
            for elem in soup.find_all(["h1", "h2", "h3", "h4", "p", "table"]):
                if elem.name in ["h1", "h2", "h3", "h4"]:
                    level = "#" * (int(elem.name[1]) + 1)
                    lines.append(f"\n{level} {elem.get_text(strip=True)}\n")
                elif elem.name == "p":
                    t = elem.get_text(strip=True)
                    if t:
                        lines.append(f"{t}\n")
                elif elem.name == "table":
                    rows = elem.find_all("tr")
                    for i, r in enumerate(rows):
                        cols = [c.get_text(strip=True) for c in r.find_all(["td", "th"])]
                        if cols:
                            lines.append("| " + " | ".join(cols) + " |")
                            if i == 0:
                                lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
                    lines.append("\n")
        else:
            lines.append(scraped_data.get("text_content", ""))

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        logger.info(f"Exported Markdown: {file_path}")

    @staticmethod
    def export_docx(scraped_data: Dict[str, Any], file_path: str):
        meta = scraped_data.get("metadata", {})
        title = meta.get("title", "Văn bản pháp luật")
        
        doc = Document()
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("THÔNG TIN VĂN BẢN", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Số hiệu: ").bold = True
        p.add_run(f"{meta.get('so_hieu', 'N/A')}\n")
        p.add_run(f"Loại văn bản: ").bold = True
        p.add_run(f"{meta.get('loai_van_ban', 'N/A')}\n")
        p.add_run(f"Ngày ban hành: ").bold = True
        p.add_run(f"{meta.get('ngay_ban_hanh', 'N/A')}\n")
        p.add_run(f"Ngày hiệu lực: ").bold = True
        p.add_run(f"{meta.get('ngay_hieu_luc', 'N/A')}\n")
        p.add_run(f"Cơ quan ban hành: ").bold = True
        p.add_run(f"{meta.get('co_quan_ban_hanh', 'N/A')}\n")
        p.add_run(f"Tình trạng: ").bold = True
        p.add_run(f"{meta.get('tinh_trang', 'N/A')}\n")
        
        doc.add_heading("NỘI DUNG VĂN BẢN", level=1)
        
        html_content = scraped_data.get("html_content", "")
        if html_content:
            soup = BeautifulSoup(html_content, "lxml")
            for elem in soup.find_all(["h1", "h2", "h3", "h4", "p", "table"]):
                if elem.name in ["h1", "h2", "h3", "h4"]:
                    doc.add_heading(elem.get_text(strip=True), level=int(elem.name[1]))
                elif elem.name == "p":
                    t = elem.get_text(strip=True)
                    if t:
                        doc.add_paragraph(t)
                elif elem.name == "table":
                    rows = elem.find_all("tr")
                    if rows:
                        first_cols = rows[0].find_all(["td", "th"])
                        table_w = len(first_cols)
                        if table_w > 0:
                            t_docx = doc.add_table(rows=len(rows), cols=table_w)
                            t_docx.style = 'Table Grid'
                            for r_idx, r in enumerate(rows):
                                cols = r.find_all(["td", "th"])
                                for c_idx, c in enumerate(cols[:table_w]):
                                    t_docx.cell(r_idx, c_idx).text = c.get_text(strip=True)
        else:
            doc.add_paragraph(scraped_data.get("text_content", ""))

        doc.save(file_path)
        logger.info(f"Exported DOCX: {file_path}")

    @staticmethod
    def _get_slug(metadata: Dict[str, Any], url: str = "") -> str:
        so_hieu = metadata.get("so_hieu", "")
        if so_hieu:
            slug = so_hieu.replace("/", "-").replace(" ", "-")
        elif url:
            # Extract slug from URL filename path
            parts = [p for p in url.split("/") if p]
            last_part = parts[-1] if parts else "van-ban"
            slug = last_part.replace(".aspx", "").replace(".html", "")
            match = re.search(r"([A-Za-z0-9\-]+-\d+)", slug)
            if match:
                slug = match.group(1)
        else:
            slug = metadata.get("title", "van-ban").replace(" ", "-")
            
        slug = re.sub(r'[\\/*?:"<>|]', "", slug)
        return slug.strip("-")[:80]
